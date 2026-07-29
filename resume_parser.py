import io
import os
import zipfile
import pypdf
import docx

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text and embedded hyperlinks from PDF file bytes."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    found_urls = set()

    for page_num, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
        
        # Extract embedded hyperlink annotations (/Annots)
        try:
            if "/Annots" in page:
                for annot in page["/Annots"]:
                    obj = annot.get_object()
                    if obj and "/A" in obj and "/URI" in obj["/A"]:
                        uri = str(obj["/A"]["/URI"]).strip()
                        if uri and uri not in found_urls:
                            found_urls.add(uri)
                            text_parts.append(f"[Hyperlink Annotation: {uri}]")
        except Exception:
            pass

    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes."""
    doc = docx.Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from TXT or raw text bytes."""
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore").strip()


def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    """
    Parses resume content from PDF, DOCX, or TXT file bytes.
    Returns cleaned string or raises ValueError if format is unsupported or unparseable.
    """
    if not file_bytes:
        raise ValueError("Uploaded resume file is empty.")

    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        extracted = extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        try:
            extracted = extract_text_from_docx(file_bytes)
        except Exception:
            # Fallback to plain text if DOCX parsing fails
            extracted = extract_text_from_txt(file_bytes)
    elif ext in [".txt", ".md", ".rtf", ".csv", ".json"]:
        extracted = extract_text_from_txt(file_bytes)
    else:
        # Fallback attempt as text
        extracted = extract_text_from_txt(file_bytes)

    if not extracted or len(extracted.strip()) < 20:
        raise ValueError("Could not extract meaningful text from resume file. Please ensure the file contains readable text.")

    # Limit to reasonable token count (~3000 words max) to keep LLM context clean
    words = extracted.split()
    if len(words) > 3000:
        extracted = " ".join(words[:3000])

    return extracted


def extract_resumes_from_zip(zip_bytes: bytes) -> list:
    """
    Unpacks a ZIP archive in memory and extracts candidate resume texts from all contained PDF/DOCX/TXT files.
    Returns list of dicts: [{"filename": str, "text": str}]
    """
    if not zip_bytes:
        raise ValueError("ZIP archive file is empty.")

    extracted_files = []
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            for info in zf.infolist():
                # Skip directories and hidden / OS metadata files
                if info.is_dir():
                    continue
                filename = os.path.basename(info.filename)
                if not filename or filename.startswith(".") or filename.startswith("__MACOSX"):
                    continue

                ext = os.path.splitext(filename)[1].lower()
                if ext in [".pdf", ".docx", ".doc", ".txt", ".md"]:
                    try:
                        file_content = zf.read(info)
                        resume_text = extract_resume_text(file_content, filename)
                        if resume_text:
                            extracted_files.append({
                                "filename": filename,
                                "text": resume_text
                            })
                    except Exception as e:
                        print(f"Skipping zip entry '{filename}' due to parse error: {e}")
    except Exception as e:
        raise ValueError(f"Invalid or corrupted ZIP archive: {str(e)}")

    if not extracted_files:
        raise ValueError("No valid PDF, DOCX, or TXT resume files were found inside the uploaded ZIP archive.")

    return extracted_files
